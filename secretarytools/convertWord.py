import argparse
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
from datetime import datetime
import os
import sys

# 民國年對照
ROC_YEAR_OFFSET = 1911

# 星期對照
WEEKDAY_ZH = ['星期一', '星期二', '星期三', '星期四', '星期五', '星期六', '星期日']


def format_date_cell(date_str, time_str, room_str):
    """
    將參數轉換成 Row 0 Col 0 的文字內容
    """
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    roc_year = dt.year - ROC_YEAR_OFFSET
    month = dt.month
    day = dt.day
    weekday = WEEKDAY_ZH[dt.weekday()]

    # 時間格式：19:00 → 晚上07:00，其他時段自動判斷
    hour, minute = map(int, time_str.split(':'))
    if 18 <= hour <= 23:
        period = "晚上"
        display_hour = hour - 12 if hour > 12 else hour
    elif 12 <= hour < 18:
        period = "下午"
        display_hour = hour - 12
    else:
        period = "早上"
        display_hour = hour

    time_display = f"{period}{display_hour:02d}:{minute:02d}"

    text = (
        f"日期：{roc_year} 年{month:02d}月{day:02d}日（{weekday}）\n"
        f"時間：{time_display}\n"
        f"地點：{room_str}"
    )
    return text


def replace_text_in_paragraph(paragraph, old_text, new_text, font_size=None):
    full_text = "".join(run.text for run in paragraph.runs)
    if old_text not in full_text:
        return False
    for i, run in enumerate(paragraph.runs):
        run.text = "" if i > 0 else full_text.replace(old_text, new_text)
        if i == 0 and font_size:
            run.font.name = "標楷體"
            run.font.size = Pt(font_size)
            if run._element.rPr is not None and hasattr(run._element.rPr, 'rFonts'):
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    return True


def replace_in_document(doc, old_text, new_text, font_size=None):
    count = 0
    for para in doc.paragraphs:
        if replace_text_in_paragraph(para, old_text, new_text, font_size):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_text_in_paragraph(para, old_text, new_text, font_size):
                        count += 1
    return count


def style_text_in_paragraph(paragraph, search_text):
    """
    尋找段落內的特定文字，並將其設為 粗體 + 斜體 (保留原段落字型)
    """
    if search_text not in paragraph.text:
        return False

    # 備份第一個 run 的字型設定，以盡量保留原格式
    font_name = None
    font_size = None
    if paragraph.runs:
        font_name = paragraph.runs[0].font.name
        font_size = paragraph.runs[0].font.size

    full_text = paragraph.text
    paragraph.clear()  # 清除段落中原有的 runs

    # 以目標字串將整段文字切開，分段重建
    parts = full_text.split(search_text)
    for i, part in enumerate(parts):
        # 加入非目標文字（維持正常格式）
        if part:
            run = paragraph.add_run(part)
            if font_name:
                run.font.name = font_name
                if run._element.rPr is not None and hasattr(run._element.rPr, 'rFonts'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if font_size:
                run.font.size = font_size
        
        # 加入目標文字（加上 粗體 + 斜體）
        if i < len(parts) - 1:
            run = paragraph.add_run(search_text)
            run.bold = True
            run.italic = True
            if font_name:
                run.font.name = font_name
                if run._element.rPr is not None and hasattr(run._element.rPr, 'rFonts'):
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if font_size:
                run.font.size = font_size
                
    return True


def delete_textbox_containing(doc, search_text):
    WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    deleted = 0
    body = doc.element.body

    for txbx in body.findall('.//' + '{' + WPS + '}txbx'):
        text = ''.join(t.text or '' for t in txbx.findall('.//' + '{' + W + '}t'))
        if search_text in text:
            node = txbx
            for _ in range(8):
                node = node.getparent()
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
                deleted += 1
                print(f"✅ 已刪除包含「{search_text}」的文字框")

    if deleted == 0:
        print(f"⚠️  未找到包含「{search_text}」的文字框，略過刪除")

    return deleted

def set_cell_text_table_align(cell, text, font_size=11):
    # 如果儲存格內已經有段落，保留第一個段落的「格式(如縮排)」，只清空文字
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.text = "" 
        run = p.add_run(text)
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        run.font.size = Pt(font_size)
        
        # 如果有多餘的段落(換行)，將其刪除以保持整潔
        for i in range(len(cell.paragraphs)-1, 0, -1):
            p_elem = cell.paragraphs[i]._element
            p_elem.getparent().remove(p_elem)
    else:
        # 如果是完全空的，才直接賦值
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "標楷體"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                run.font.size = Pt(font_size)

def set_cell_text(cell, text, font_size=11):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "標楷體"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            run.font.size = Pt(font_size)


def modify_word_file(input_path, date_str=None, time_str=None, room_str=None, highlight_text=None):
    if not os.path.exists(input_path):
        print(f"❌ 找不到檔案：{input_path}")
        return False

    base_name = os.path.splitext(input_path)[0]
    output_path = f"{base_name}_修改版.docx"

    doc = Document(input_path)
    modified = False

    if not doc.tables:
        print("⚠️  檔案中沒有找到表格")
        return False

    table = doc.tables[0]

    # ====================== 1. Row 0, Column 0 (日期區塊) ======================
    if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
        cell = table.rows[0].cells[0]
        if date_str and time_str and room_str:
            try:
                cell_text = format_date_cell(date_str, time_str, room_str)
            except ValueError as e:
                print(f"❌ 日期/時間格式錯誤：{e}")
                return False
        else:
            print(f"❌ 缺少日期/時間/地點參數")
            return False
            
        set_cell_text(cell, cell_text)
        print(f"✅ Row 0, Column 0 已修改（日期區塊）：\n   {cell_text.replace(chr(10), ' / ')}")
        modified = True

    # ====================== 2. Row 0, Column 2 (主席/司儀/電腦) ======================
    if len(table.rows) > 0 and len(table.rows[0].cells) > 2:
        cell = table.rows[0].cells[2]
        set_cell_text(cell, "主席：李自強會長 擔任主席\n司儀：楊國華\n電腦：辜女育")
        print("✅ Row 0, Column 2 已修改（主席 / 司儀 / 電腦）")
        modified = True

    # ====================== 3. Row 3, Column 0 ======================
    if len(table.rows) > 3 and len(table.rows[3].cells) > 0:
        cell = table.rows[3].cells[0]
        set_cell_text_table_align(cell, "9.宣講員宣講")
        print("✅ Row 3, Column 0 已修改 → 9. 宣講員宣講")
        modified = True

    # ====================== 4. Row 4, Column 4 ======================
    if len(table.rows) > 4 and len(table.rows[4].cells) > 4:
        cell = table.rows[4].cells[4]
        set_cell_text_table_align(cell, "16.輔導法師開示")
        print("✅ Row 4, Column 4 已修改 → 16.輔導法師開示")
        modified = True

    # ====================== 5. 全文替換特定字串 (動態讀取 highlight_text) ======================
    if highlight_text:
        new_title = f"{highlight_text}月例會活動討論案"
        replace_count = replace_in_document(
            doc,
            old_text="聯合月例會活動討論案",
            new_text=new_title,
            font_size=20
        )
        if replace_count > 0:
            print(f"✅ 已將「聯合月例會活動討論案」替換為「{new_title}」（共 {replace_count} 處，字體 20pt）")
            modified = True
        else:
            print("⚠️  未找到「聯合月例會活動討論案」，略過替換")
    else:
        print("⚠️  未提供 -n 參數，略過替換「聯合月例會活動討論案」")

    # ====================== 6. 針對表格特定文字加上 粗體 + 斜體 ======================
    if highlight_text:
        highlight_count = 0
        for tb in doc.tables:
            for row in tb.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if style_text_in_paragraph(para, highlight_text):
                            highlight_count += 1
        
        if highlight_count > 0:
            print(f"✅ 已將表格中所有的「{highlight_text}」設為 粗體 + 斜體 (共 {highlight_count} 個段落)")
            modified = True
        else:
            print(f"⚠️  未在表格中找到「{highlight_text}」，無法設定粗斜體")

    # ====================== 7. 刪除含「聯合月例會】」的文字框 ======================
    del_count = delete_textbox_containing(doc, "聯合月例會】")
    if del_count > 0:
        modified = True

    # ================== 儲存檔案 ==================
    if modified:
        doc.save(output_path)
        print(f"\n🎉 修改完成！")
        print(f"   修改後檔案已儲存為：\n   {output_path}")
    else:
        print("\n⚠️  修改未成功，請檢查文件與參數。")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="修改月例會 Word 檔案")
    parser.add_argument("-f", "--file",  required=True, help="Word 檔案路徑")
    parser.add_argument("-d", "--date",  help="日期，格式 YYYY-MM-DD，例如 2026-04-08")
    parser.add_argument("-t", "--time",  help="時間，格式 HH:MM，例如 19:00")
    parser.add_argument("-m", "--room",  help="房號，例如 305")
    parser.add_argument("-n", "--name",  help="分會名稱 (如: 宜六)，將用於替換標題及在表格中設為粗斜體")
    args = parser.parse_args()

    # d, t, m 三個參數必須同時提供或同時不提供
    date_time_room = [args.date, args.time, args.room]
    if any(date_time_room) and not all(date_time_room):
        print("❌ -d、-t、-m 三個參數必須同時提供")
        sys.exit(1)

    modify_word_file(
        input_path=args.file,
        date_str=args.date,
        time_str=args.time,
        room_str=args.room,
        highlight_text=args.name  # 傳入 -n 的值
    )